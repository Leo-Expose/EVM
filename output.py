import serial
import time
import os
from flask import Flask, jsonify
from threading import Thread
import tkinter as tk
from tkinter import messagebox
from docx import Document

# Define the serial port and baud rate.
SERIAL_PORT = 'COM4'  # Update this to the correct port
BAUD_RATE = 9600
BACKUP_FILE = os.path.expanduser("~/Desktop/Voting/votes_backup.txt")

# Candidate names
candidates = ["A", "B", "C", "X", "Y", "Z"]
votes = {candidate: 0 for candidate in candidates}

# Load existing votes from backup file
if os.path.exists(BACKUP_FILE):
    with open(BACKUP_FILE, 'r') as f:
        for line in f:
            name, vote_count = line.strip().split(',')
            votes[name] = int(vote_count)

# Initialize serial connection
ser = serial.Serial(SERIAL_PORT, BAUD_RATE)

def save_backup():
    with open(BACKUP_FILE, 'w') as f:
        for name, count in votes.items():
            f.write(f"{name},{count}\n")

def export_results():
    result_file = os.path.expanduser("~/Desktop/Voting/FINAL.docx")
    doc = Document()
    doc.add_heading('Voting Results', 0)
    for name, count in votes.items():
        doc.add_paragraph(f"{name}: {count} votes")
    doc.save(result_file)
    print("Results saved to", result_file)

def print_results():
    result_file = os.path.expanduser("~/Desktop/Voting/results.txt")
    with open(result_file, 'w') as f:
        for name, count in votes.items():
            f.write(f"{name}: {count} votes\n")
    print("Results saved to", result_file)

def stop_voting():
    global running
    running = False
    messagebox.showinfo("Voting Stopped", "Voting has been stopped and results exported.")
    export_results()
    root.destroy()

# Flask app setup
app = Flask(__name__)

@app.route('/votes')
def get_votes():
    return jsonify(votes)

def start_flask():
    app.run(debug=True, use_reloader=False, host='192.168.1.1', port=5000)

def start_voting():
    global running
    running = True
    try:
        while running:
            if ser.in_waiting > 0:
                vote = ser.readline().decode().strip()
                if vote.startswith("Vote: "):
                    candidate_index = int(vote.split(" ")[1])
                    candidate_name = candidates[candidate_index]
                    votes[candidate_name] += 1
                    print(f"{candidate_name} received a vote. Total now: {votes[candidate_name]}")
                    save_backup()
            time.sleep(0.1)
    except KeyboardInterrupt:
        print("Vote counting stopped.")
        print_results()
    finally:
        ser.close()

# Start Flask server in a separate thread
flask_thread = Thread(target=start_flask)
flask_thread.start()

# Tkinter GUI setup
root = tk.Tk()
root.title("Voting Control")

stop_button = tk.Button(root, text="Stop Voting", command=stop_voting, font=('Helvetica', 16))
stop_button.pack(pady=20)

root.geometry("300x200")
root.protocol("WM_DELETE_WINDOW", stop_voting)

# Start the voting process
voting_thread = Thread(target=start_voting)
voting_thread.start()

# Start the Tkinter main loop
root.mainloop()
